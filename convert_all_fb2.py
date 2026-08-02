import os, re, sys, json, xml.etree.ElementTree as ET, time
from docx import Document
from docx.shared import Pt
import win32com.client

TEMP_DIR = r"D:\AI_Project\temp_docx"
os.makedirs(TEMP_DIR, exist_ok=True)

def extract_fb2_text(path):
    tree = ET.parse(path)
    root = tree.getroot()
    ns = {"fb": "http://www.gribuser.ru/xml/fictionbook/2.0"}
    texts = []
    for body in root.findall(".//fb:body", ns):
        for p in body.findall(".//fb:p", ns):
            if p.text:
                texts.append(p.text)
    return "\n".join(texts)

def fb2_to_docx(fb2_path, docx_path):
    text = extract_fb2_text(fb2_path)
    doc = Document()
    for para in text.split("\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        run = p.add_run(para[:5000])
        run.font.size = Pt(10)
    doc.save(docx_path)

# Find all FB2 without PDF
root_dir = r"D:\AI_Project\Книги"
to_convert = []
for dirpath, dirnames, filenames in os.walk(root_dir):
    dirnames[:] = [d for d in dirnames if not d.startswith('.') and not d.startswith('__') and d not in ('converted', 'converted_new', 'outputs', 'summary', 'temp_docx')]
    for f in filenames:
        if not f.lower().endswith('.fb2'):
            continue
        stem = os.path.splitext(f)[0]
        pdf_path = os.path.join(dirpath, stem + ".pdf")
        if not os.path.exists(pdf_path):
            to_convert.append((dirpath, f, os.path.join(dirpath, f)))

print(f"Total FB2 to convert: {len(to_convert)}")

# Phase 1: FB2 -> DOCX
docx_files = []
for dirpath, fname, fpath in to_convert:
    stem = os.path.splitext(fname)[0]
    safe_name = re.sub(r'[^\w\-_. ]', '_', stem)[:80]
    docx_path = os.path.join(TEMP_DIR, safe_name + ".docx")
    try:
        fb2_to_docx(fpath, docx_path)
        docx_files.append((dirpath, stem, docx_path))
        size = os.path.getsize(docx_path)
        print(f"  DOCX OK: {fname[:50]:50s} -> {size/1024:.0f} KB")
    except Exception as e:
        print(f"  DOCX FAIL: {fname[:50]} -> {str(e)[:60]}")

print(f"\nDOCX created: {len(docx_files)}. Phase 2: DOCX -> PDF via Word...")

# Phase 2: DOCX -> PDF (batch via Word)
word = win32com.client.Dispatch("Word.Application")
word.Visible = False

batch_size = 20
for i in range(0, len(docx_files), batch_size):
    batch = docx_files[i:i+batch_size]
    print(f"\nBatch {i//batch_size + 1}/{(len(docx_files)-1)//batch_size + 1} ({len(batch)} files):")
    for dirpath, stem, docx_path in batch:
        pdf_path = os.path.join(dirpath, stem + ".pdf")
        try:
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close()
            size = os.path.getsize(pdf_path)
            print(f"  PDF OK: {stem[:50]:50s} -> {size/1024:.0f} KB")
        except Exception as e:
            print(f"  PDF FAIL: {stem[:50]} -> {str(e)[:60]}")
    # Restart Word between batches to avoid memory issues
    if i + batch_size < len(docx_files):
        print("  Restarting Word...")
        word.Quit()
        time.sleep(2)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

word.Quit()
print(f"\nDone! Converted {len(docx_files)} FB2 to PDF")

# Cleanup temp
import shutil
try:
    shutil.rmtree(TEMP_DIR)
    print("Temp files cleaned")
except:
    pass
