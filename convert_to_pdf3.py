import os, sys, re, zipfile, xml.etree.ElementTree as ET, subprocess, time
from docx import Document
from docx.shared import Pt

def extract_fb2(path):
    tree = ET.parse(path)
    root = tree.getroot()
    ns = {"fb": "http://www.gribuser.ru/xml/fictionbook/2.0"}
    texts = []
    for body in root.findall(".//fb:body", ns):
        for p in body.findall(".//fb:p", ns):
            if p.text:
                texts.append(p.text)
    return "\n".join(texts)

def extract_epub(path):
    text = ""
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if name.endswith((".htm", ".html", ".xhtml")):
                content = z.read(name).decode("utf-8", errors="replace")
                clean = re.sub(r"<[^>]+>", "", content)
                clean = re.sub(r"\s+", " ", clean)
                text += clean + "\n"
    return text

def txt_to_docx(text, docx_path, max_para_len=5000):
    doc = Document()
    for para in text.split("\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        run = p.add_run(para[:max_para_len])
        run.font.size = Pt(10)
    doc.save(docx_path)

def fb2_to_docx(fb2_path, docx_path):
    text = extract_fb2(fb2_path)
    txt_to_docx(text, docx_path)

def epub_to_docx(epub_path, docx_path):
    text = extract_epub(epub_path)
    txt_to_docx(text, docx_path)

def docx_to_pdf(docx_path, pdf_path):
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path))
        try:
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        finally:
            doc.Close()
    finally:
        word.Quit()

# === CONFIG ===
FOLDERS = [
    r"D:\AI_Project\Книги\свободные_книги",
    r"D:\AI_Project\Книги\Развитие",
    r"D:\AI_Project\health\books",
]

SKIP_PREFIXES = ("анализ_", "анализ ", "общий ")
SKIP_NAMES = {"Список нужных книг.txt", "psixologiia-stressa.txt"}
TARGET_EXTS = {".fb2", ".epub", ".txt"}

TEMP_DIR = r"D:\AI_Project\converted_new"
os.makedirs(TEMP_DIR, exist_ok=True)

total = 0
ok = 0
fail = 0

for folder in FOLDERS:
    if not os.path.exists(folder):
        continue
    for fname in sorted(os.listdir(folder)):
        fpath = os.path.join(folder, fname)
        if os.path.isdir(fpath):
            continue
        ext = os.path.splitext(fname)[1].lower()
        if ext not in TARGET_EXTS:
            continue
        if any(fname.startswith(p) for p in SKIP_PREFIXES):
            continue
        if fname in SKIP_NAMES:
            continue

        stem = os.path.splitext(fname)[0]
        pdf_path = os.path.join(folder, stem + ".pdf")
        if os.path.exists(pdf_path):
            continue

        total += 1
        safe_stem = re.sub(r'[^\w\-_. ]', '_', stem)[:60]
        docx_path = os.path.join(TEMP_DIR, safe_stem + ".docx")

        print(f"[{total}] {fname[:55]:55s}...", end=" ")
        try:
            # Step 1: extract text and create docx
            if ext == ".fb2":
                txt_to_docx(extract_fb2(fpath), docx_path)
            elif ext == ".epub":
                txt_to_docx(extract_epub(fpath), docx_path)
            elif ext == ".txt":
                text = open(fpath, "r", encoding="utf-8", errors="replace").read()
                if len(text) < 50:
                    raise Exception("Too short")
                txt_to_docx(text, docx_path)

            size_docx = os.path.getsize(docx_path)
            if size_docx < 100:
                raise Exception(f"DOCX too small: {size_docx} bytes")

            # Step 2: convert docx to pdf via Word
            docx_to_pdf(docx_path, pdf_path)
            size_pdf = os.path.getsize(pdf_path)
            print(f"OK ({size_pdf/1024:.0f} KB)")
            ok += 1

        except Exception as e:
            print(f"FAIL: {str(e)[:80]}")
            fail += 1

        # Cleanup temp docx
        try:
            if os.path.exists(docx_path):
                os.remove(docx_path)
        except:
            pass

print(f"\nDone: {ok} OK, {fail} FAIL out of {total}")
